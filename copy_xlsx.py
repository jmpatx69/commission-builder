import streamlit as st
import pandas as pd
from io import BytesIO
import zipfile
import openpyxl
from docx import Document
import os
import locale

# Set the locale to the user's default setting (for number formatting)
locale.setlocale(locale.LC_ALL, '')

def excel_col_to_index(col_letter):
    col_letter = col_letter.upper().strip()
    index = 0
    for char in col_letter:
        index = index * 26 + (ord(char) - ord('A') + 1)
    return index - 1

def replace_in_word_doc(doc, replacements):
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    # Replace in headers and footers
    for section in doc.sections:
        for header in section.header.paragraphs:
            for key, value in replacements.items():
                if key in header.text:
                    header.text = header.text.replace(key, str(value))
        for footer in section.footer.paragraphs:
            for key, value in replacements.items():
                if key in footer.text:
                    footer.text = footer.text.replace(key, str(value))

    return doc

st.title("Commission Doc Generator")

uploaded_sales_file = st.file_uploader("Upload Sales Data (XLSX)", type=["xlsx"])
template_folder = st.text_input("Enter Path to Workbook Template Folder (e.g., /path/to/folder)")
word_template_folder = st.text_input("Enter Path to Word Doc Templates Folder (e.g., /path/to/folder)")

if uploaded_sales_file and template_folder and word_template_folder:
    sales_data = pd.ExcelFile(uploaded_sales_file)
    
    header_row = None
    for i in range(50):
        df_temp = sales_data.parse("Target Summary", header=i)
        if "Sales Person" in df_temp.columns:
            header_row = i
            break
    
    if header_row is None:
        st.error("The input file must contain a 'Sales Person' column within the first 50 rows.")
    else:
        df = sales_data.parse("Target Summary", header=header_row)
        required_cols = ["Sales Person", "Territory", "Workbook Template", "Doc Template"]
        if not all(col in df.columns for col in required_cols):
            st.error("The input file must contain 'Sales Person', 'Territory', 'Workbook Template', and 'Doc Template' columns.")
        else:
            sales_reps = df["Sales Person"].unique()
            output_buffers = []

            for rep in sales_reps:
                rep_data = df[df["Sales Person"] == rep]
                if rep_data.empty:
                    continue

                # Get the workbook template for the current sales rep
                workbook_template_name = rep_data["Workbook Template"].iloc[0]
                workbook_template_path = os.path.join(template_folder, workbook_template_name)
                
                if not os.path.exists(workbook_template_path):
                    st.warning(f"Workbook template '{workbook_template_name}' not found in '{template_folder}' for {rep}.")
                    continue

                # Workbook Processing
                output_wb = openpyxl.load_workbook(workbook_template_path)
                if "Component Input" not in output_wb.sheetnames:
                    st.error(f"The workbook template '{workbook_template_name}' must contain a 'Component Input' tab.")
                    continue

                output_ws = output_wb["Component Input"]
                component_input = pd.read_excel(workbook_template_path, sheet_name="Component Input")
                workbook_mapping = component_input.iloc[:, 5].fillna('').tolist()
                word_doc_mapping = component_input.iloc[:, 6].fillna('').tolist()

                for row_idx, col_letter in enumerate(workbook_mapping):
                    if col_letter:
                        source_col_idx = excel_col_to_index(col_letter)
                        target_col_idx = 3
                        if source_col_idx < len(rep_data.columns):
                            value_to_copy = rep_data.iloc[0, source_col_idx]
                            output_ws.cell(row=row_idx + 2, column=target_col_idx + 1, value=value_to_copy)

                output_filename = f"{rep}_Sales_Report.xlsx"
                output_buffer = BytesIO()
                output_wb.save(output_buffer)
                output_buffer.seek(0)
                output_buffers.append((output_filename, output_buffer))

                # Word Doc Processing
                word_doc_filename = rep_data["Doc Template"].iloc[0]
                word_doc_path = os.path.join(word_template_folder, word_doc_filename)
                
                if os.path.exists(word_doc_path):
                    doc = Document(word_doc_path)
                    replacements = {
                        "<<NAME>>": rep,
                        "<<TERR>>": rep_data["Territory"].iloc[0]
                    }

                    for row_idx, mapping in enumerate(word_doc_mapping):
                        if mapping and (mapping.startswith("Q") or mapping.startswith("A")):
                            source_col_idx = excel_col_to_index(workbook_mapping[row_idx])
                            if source_col_idx < len(rep_data.columns):
                                value = rep_data.iloc[0, source_col_idx]
                                if pd.api.types.is_numeric_dtype(type(value)):
                                    value = "${:,.0f}".format(int(value)) if value % 1 == 0 else "${:,.2f}".format(value)
                                replacements[f"<<{mapping}>>"] = value if value is not None else "N/A"

                    updated_doc = replace_in_word_doc(doc, replacements)
                    word_output_filename = word_doc_filename.replace("Master.docx", f"_{rep}.docx")
                    word_output_buffer = BytesIO()
                    updated_doc.save(word_output_buffer)
                    word_output_buffer.seek(0)
                    output_buffers.append((word_output_filename, word_output_buffer))
                else:
                    st.warning(f"Word doc template '{word_doc_filename}' not found in '{word_template_folder}' for {rep}.")

            # Create ZIP
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                for filename, buffer in output_buffers:
                    zip_file.writestr(filename, buffer.getvalue())
            zip_buffer.seek(0)

            zip_filename = "Commission_Workbooks.zip"

            st.success(f"Generated {len(output_buffers)} files!")
            st.download_button(
                label="Download All Sales Reports (ZIP)",
                data=zip_buffer,
                file_name=zip_filename,
                mime="application/zip"
            )